"""Helper that writes the new generate_proposal_docx.py"""
script = r'''
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x00, 0x3F, 0x87)
GOLD  = RGBColor(0xC8, 0x96, 0x3E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x2C, 0x2C, 0x2C)
MID   = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x88, 0x88, 0x88)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

def shade_para(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def shade_cell(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def cell_border(cell, color='003F87', sz='6'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcB = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement('w:' + edge)
        tag.set(qn('w:val'), 'single'); tag.set(qn('w:sz'), sz)
        tag.set(qn('w:space'), '0'); tag.set(qn('w:color'), color)
        tcB.append(tag)
    tcPr.append(tcB)

def add_run(para, text, bold=False, italic=False, size=11, color=None, font='Cambria'):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic; r.font.name = font; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r

def body(text, bold_pre=None, color=MID, size=11, sb=6, sa=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if bold_pre: add_run(p, bold_pre + ' ', bold=True, size=size, color=DARK)
    add_run(p, text, size=size, color=color)
    return p

def bullet(text, bold_pre=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    if bold_pre: add_run(p, bold_pre + ' ', bold=True, size=11, color=DARK)
    add_run(p, text, size=11, color=MID)
    return p

def callout(text, bold_pre=None, gold=False):
    fill = 'FDF6E8' if gold else 'E8F0FA'; bcol = 'C8963E' if gold else '003F87'
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    shade_para(p, fill)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '14')
    left.set(qn('w:space'), '4'); left.set(qn('w:color'), bcol)
    pBdr.append(left); pPr.append(pBdr)
    if bold_pre: add_run(p, bold_pre + ' ', bold=True, size=11, color=GOLD if gold else NAVY)
    add_run(p, text, size=11, color=DARK)

def sec_hdr(num, eyebrow, title):
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(16); p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(num + '  '); r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(26); r1.font.color.rgb = NAVY
    r2 = p1.add_run(eyebrow.upper()); r2.font.name = 'Arial'; r2.font.size = Pt(9); r2.font.color.rgb = GOLD
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(6)
    rt = p2.add_run(title); rt.bold = True; rt.font.name = 'Arial'; rt.font.size = Pt(20); rt.font.color.rgb = NAVY
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0); rule.paragraph_format.space_after = Pt(12)
    pPr = rule._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '003F87')
    pBdr.append(bot); pPr.append(pBdr)

def subhead(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text.upper()); r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(11); r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'D0D9E8')
    pBdr.append(bot); pPr.append(pBdr)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'D0D9E8')
    pBdr.append(bot); pPr.append(pBdr)

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade_cell(c, '003F87')
        p = c.paragraphs[0]; rr = p.add_run(h)
        rr.bold = True; rr.font.name = 'Arial'; rr.font.size = Pt(10); rr.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    for ri, row_data in enumerate(rows):
        fill = 'F5F8FD' if ri % 2 == 0 else 'FFFFFF'
        tr = t.rows[ri+1]
        for ci, txt in enumerate(row_data):
            c = tr.cells[ci]; shade_cell(c, fill); p = c.paragraphs[0]
            is_gold = txt.startswith('>>')
            display = txt[2:] if is_gold else txt
            rr = p.add_run(display); rr.font.name = 'Cambria'; rr.font.size = Pt(10)
            if is_gold: rr.bold = True; rr.font.color.rgb = GOLD
            else: rr.font.color.rgb = DARK
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows: r.cells[i].width = Inches(w)
    return t

# ── COVER ──────────────────────────────────────────────────
p = doc.add_paragraph(); shade_para(p, '003F87')
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.add_run(' ')
p = doc.add_paragraph(); shade_para(p, 'C8963E')
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.add_run(' ')
doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30); p.paragraph_format.space_after = Pt(4)
add_run(p, 'Pithonix', bold=True, size=30, color=NAVY, font='Arial')
add_run(p, 'AI', bold=True, size=30, color=GOLD, font='Arial')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
add_run(p, "INDIA'S AI-NATIVE ENTERPRISE BUILDER", size=8, color=LIGHT, font='Arial')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(2)
add_run(p, 'Submitted to', size=9, color=LIGHT, font='Arial')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
add_run(p, 'SOFTWARE TECHNOLOGY PARKS OF INDIA', bold=True, size=14, color=NAVY, font='Arial')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(24)
add_run(p, 'Ministry of Electronics & IT (MeitY)', size=10, color=MID, font='Arial')

p = doc.add_paragraph(); shade_para(p, 'C8963E')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(14); p.add_run('      ')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(14)
add_run(p, 'STRATEGIC PARTNERSHIP PROPOSAL  |  APRIL 2026', size=9, color=GOLD, font='Arial')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
add_run(p, 'GCC Playbook', size=36, color=NAVY, font='Georgia')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
add_run(p, 'Strategic Platform Partnership', italic=True, size=30, color=GOLD, font='Georgia')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(22)
add_run(p, 'Proposal', size=26, color=NAVY, font='Georgia')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(26)
add_run(p, "A proposal for co-hosting and co-branding India's first AI-native GCC Intelligence Platform on the STPI digital ecosystem, accelerating foreign investment and positioning STPI as the digital front door for India's $100B+ GCC ambition.", size=11, color=MID, font='Arial')

t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val) in enumerate([
    ('PREPARED BY', 'PITHONIX AI INDIA PRIVATE LIMITED'),
    ('DATE', 'April 2026'),
    ('PLATFORM', 'gcc-playbook.pithonix.ai'),
]):
    c = t.rows[0].cells[i]; shade_cell(c, 'F5F8FD'); cell_border(c)
    p1 = c.paragraphs[0]; add_run(p1, lbl, size=8, color=LIGHT, font='Arial')
    p1.paragraph_format.space_after = Pt(2)
    p2 = c.add_paragraph(); add_run(p2, val, bold=True, size=10, color=NAVY, font='Arial')
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(22)
add_run(p, 'PRIVATE & CONFIDENTIAL  |  FOR STPI REVIEW ONLY', bold=True, size=9, color=GOLD, font='Arial')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
add_run(p, 'Confidential  |  Prepared by PITHONIX AI INDIA PRIVATE LIMITED for STPI  |  April 2026', size=8, color=LIGHT, font='Arial')

p = doc.add_paragraph(); shade_para(p, '003F87')
p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(0); p.add_run(' ')
doc.add_page_break()

# ── TOC ───────────────────────────────────────────────────
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(20)
add_run(p, 'TABLE OF CONTENTS', size=10, color=GOLD, font='Arial')

toc = [
    ('01', 'Executive Summary', 'The GCC opportunity, the platform, and the proposal at a glance'),
    ('02', 'The Problem: Why Foreign Investors Struggle', 'Market fragmentation, consultant dependency, and investor drop-off'),
    ('03', 'The Solution: GCC Playbook Platform', 'Platform capabilities, technology stack, and live demonstration'),
    ('04', 'Why STPI: The Strategic Fit', 'Institutional alignment, geographic coverage, and competitive positioning'),
    ('05', 'Partnership Options: Three Models', 'Embed, White Label, and Co-Development with comparative analysis'),
    ('06', 'Technical Brief: Integration and Data Security', 'Integration architecture, security posture, and hosting options'),
    ('07', 'Co-Branding Proposal', 'Visual identity, STPI theme integration, and platform co-branding'),
    ('08', 'Impact Projections', 'Investor reach, GCC creation, and STPI registration uplift estimates'),
    ('09', 'Next Steps and Timeline', 'Six-week implementation roadmap from MOU to live launch'),
    ('10', 'About Pithonix AI', 'Company background, flagship products, and GCC credentials'),
]
for num, name, sub in toc:
    t = doc.add_table(rows=1, cols=2)
    c0 = t.rows[0].cells[0]; c0.width = Inches(0.55)
    p0 = c0.paragraphs[0]; add_run(p0, num, bold=True, size=15, color=GOLD, font='Arial')
    p0.paragraph_format.space_before = Pt(4); p0.paragraph_format.space_after = Pt(2)
    c1 = t.rows[0].cells[1]; p1 = c1.paragraphs[0]
    add_run(p1, name, bold=True, size=12, color=NAVY, font='Arial')
    p1.paragraph_format.space_before = Pt(4); p1.paragraph_format.space_after = Pt(2)
    p2 = c1.add_paragraph(); add_run(p2, sub, size=10, color=MID)
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(8)
doc.add_page_break()

# ── S01: EXECUTIVE SUMMARY ────────────────────────────────
sec_hdr('01', 'Section One', 'Executive Summary')
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (num, lbl) in enumerate([
    ('1,800+', 'GCCs operating in India (2025)'),
    ('$64B+', 'GCC sector annual revenue'),
    ('$100B+', 'Projected GCC revenue by 2030'),
]):
    c = t.rows[0].cells[i]; shade_cell(c, 'E8F0FA'); cell_border(c)
    p1 = c.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p1, num, bold=True, size=24, color=GOLD if i == 1 else NAVY, font='Arial')
    p1.paragraph_format.space_before = Pt(8); p1.paragraph_format.space_after = Pt(2)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, lbl, size=9, color=MID, font='Arial')
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(8)
doc.add_paragraph()
body("India's GCC sector is the fastest-growing segment of the global technology economy. 1,800+ GCCs, $64 billion in annual revenue, 1.9 million professionals, and a clear trajectory to $100 billion by 2030. No other country comes close. The world has already decided that India is where capability gets built.")
body("But here is the problem nobody is talking about. A CFO in Germany, a CTO in Singapore, a board in New York. They want to set up in India. They are motivated. They are ready. And then they spend three months navigating Invest India, Make in India, NASSCOM reports, state investment boards, STPI documentation, and commercial real estate surveys. None of it connected, none of it personalized, none of it answering the one question they actually have. They spend $50,000 to $200,000 just to decide. Mid-market companies, the next generation of GCC investors, simply give up.")
callout("that automates the entire discovery-to-blueprint process using Gemini AI. A foreign company enters their industry, headcount target, functions, and timeline, and receives a complete, personalized GCC Blueprint in under 10 seconds. Free of charge. No consultant required.",
    bold_pre="Pithonix AI has built India's first AI-native GCC Intelligence Platform, the GCC Playbook,")
body("This proposal is for STPI, India's apex IT/ITES export body under MeitY, to co-host the GCC Playbook on stpi.in. Not a complex integration. Not a procurement project. One embed: STPI becomes the first government body in India to offer foreign investors a live, AI-generated GCC blueprint at the exact moment they are looking for one.")
body("The case is straightforward. For STPI, this turns stpi.in from a registration portal into an investor acquisition engine. For Pithonix, STPI's institutional name is the most powerful trust signal in the Indian government ecosystem. For India: more GCCs, more jobs, more FDI, from companies that would otherwise have given up.")
callout("No engineering resources. No procurement cycle. No infrastructure investment. The platform can be live on stpi.in within two weeks of signing a Letter of Intent. The only input required from STPI is institutional endorsement, and the return is leadership of India's AI-powered investment promotion story.",
    bold_pre="This is a zero-cost, high-impact initiative for STPI.", gold=True)
divider()

# ── S02: THE PROBLEM ──────────────────────────────────────
sec_hdr('02', 'Section Two', 'The Problem: Why Foreign Investors Struggle')
body("India is the obvious choice for a GCC. And yet, too many foreign companies walk away before they even start. Not because India isn't ready. Because the information is not.")
bullet("A CFO in Germany evaluating a 200-person technology GCC in India has to navigate Invest India, Make in India, state investment boards, NASSCOM reports, STPI documentation, and commercial real estate portals. None connected, none personalized, none telling her what she actually needs to know. Right now, for her company, for her budget.", bold_pre="There is no single place to get a real answer.")
bullet("Foreign investors landing on stpi.in are at peak intent. They are actively evaluating India. But the portal is designed for post-decision registration, not pre-decision discovery. That gap is where opportunities are lost every week.", bold_pre="STPI's portal is built for companies that have already decided. Not for companies still deciding.")
bullet("$50,000 to $200,000. That is what a serious GCC feasibility study costs through a Big 4 or GCC-specialist advisory firm. Large MNCs pay it. Mid-market companies, PE-backed businesses, growth-stage technology firms. They cannot. And so they don't come. India loses the GCC investments it deserves most.", bold_pre="The advice market is expensive. And exclusive.")
bullet("The question every foreign investor actually asks: 'I want a 150-person data engineering GCC in India for $3 million over 3 years. Which city? Which structure? Which incentives? What do I do first?' This question goes unanswered by every government platform that exists today.", bold_pre="No portal answers the real question.")
bullet("Industry estimates, and Pithonix's own engagement data, show that investor drop-off in the first 60 to 90 days of GCC evaluation is the rule, not the exception. Information overload. Contradictory guidance. Cost of advice. India is losing investments it should be winning.", bold_pre="Two to three companies walk away for every one that stays.")
callout("The answer is not another brochure. It is not another advisory meeting. It is an intelligent platform, free, instant, and personalized, that answers the real question, for the real investor, at the exact moment they are asking it, and walks them straight into STPI's ecosystem.")
divider()
doc.add_page_break()

# ── S03: THE SOLUTION ─────────────────────────────────────
sec_hdr('03', 'Section Three', 'The Solution: GCC Playbook Platform')
body("The GCC Playbook is India's first AI-native GCC Intelligence Platform, built and operated by PITHONIX AI INDIA PRIVATE LIMITED. The platform is live, functional, and accessible at gcc-playbook.pithonix.ai. It requires no login, no consultant, and no prior knowledge of Indian regulatory frameworks. A foreign executive anywhere in the world can receive a complete, personalized GCC blueprint for India in under 10 seconds.")
subhead('Core Platform Capabilities')
bullet("The platform's flagship tool. A foreign company inputs their industry vertical, target headcount, primary functions (engineering, analytics, finance, legal operations, etc.), timeline, and budget range. Powered by Gemini AI with 2026 India market benchmarks sourced from NASSCOM, CBRE, Mercer, and Zinnov data, the simulator generates a complete, personalized GCC Blueprint in under 10 seconds, at no cost.", bold_pre="AI-Powered GCC Simulator.")
bullet("Every generated blueprint includes: recommended Indian city drawn from a database of 25 cities, a detailed 3-year cost projection in USD covering talent, real estate, setup costs and operational expenses, government incentives applicable to the user's industry and location profile, banking and treasury support options, entity formation steps under Indian company law, a Phase 0 action plan covering the first 90 days, and a risk and mitigation map for the GCC establishment journey.", bold_pre="Comprehensive Blueprint Output.")
bullet("An interactive comparative intelligence module covering 25 cities across India, including all major Tier 1 destinations (Bengaluru, Hyderabad, Mumbai, Pune, Chennai, Gurgaon, Noida) as well as emerging Tier 2 and Tier 3 locations (Coimbatore, Kochi, Jaipur, Ahmedabad, Chandigarh, Nagpur, Vizag, and others). Each city profile includes real-time talent pool data, current office rental rates, attrition rates by function, and applicable state and central government incentive schemes.", bold_pre="India City Intelligence Database.")
bullet("A structured, navigable directory of central and state government incentives available to GCC investors, with each incentive card linking directly to the relevant government portal: STPI, DPIIT, MeitY, Invest India, and individual state IT departments. STPI's STP scheme is prominently featured as a primary recommendation for eligible entities.", bold_pre="Government Incentive Cards.")
bullet("The platform's advanced reasoning engine, built on Pithonix AI's proprietary Graph of Thought (GOT) architecture, that answers free-form, open-ended GCC strategy questions. Investors can ask questions like 'How does India's talent market compare to Poland for an AI/ML GCC?' and receive structured, multi-dimensional analysis drawing on the platform's knowledge base.", bold_pre="GOT Advisor (Graph of Thought).")
subhead('Technology Stack')
tech_items = [
    ('FRONTEND',      'Single-page HTML/JavaScript, zero framework lock-in. Deployable on any government-approved infrastructure including NIC Cloud and MeghRaj without modification.'),
    ('AI ENGINE',     "Google Gemini 2.5 Flash, accessed via a secure server-side proxy. The API key is never exposed to the browser. All prompts include 2026 India market benchmark data."),
    ('HOSTING',       "Vercel global CDN (current). 99.99% uptime SLA. Can be mirrored on NIC Cloud or MeghRaj within 4 weeks if required by STPI's data residency policy."),
    ('LEAD CAPTURE',  "Microsoft Power Automate pipeline routing investor inquiries to Pithonix CRM and STPI's preferred notification endpoint in real time."),
    ('DATA SOURCES',  '2026 benchmarks: NASSCOM GCC Report, CBRE India Office Market, Mercer Total Remuneration Survey, Zinnov GCC Landscape Report. Updated quarterly.'),
    ('ACCESSIBILITY', 'WCAG 2.1 AA compliant. Mobile-responsive. Tested on 3G connections. Supports screen readers. Available in English; Hindi and regional language versions in roadmap.'),
]
t = doc.add_table(rows=len(tech_items), cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (lbl, desc) in enumerate(tech_items):
    fill = 'E8F0FA' if i % 2 == 0 else 'F5F8FD'
    c0 = t.rows[i].cells[0]; c0.width = Inches(1.3); shade_cell(c0, fill)
    p0 = c0.paragraphs[0]; add_run(p0, lbl, bold=True, size=9, color=NAVY, font='Arial')
    p0.paragraph_format.space_before = Pt(5); p0.paragraph_format.space_after = Pt(5)
    c1 = t.rows[i].cells[1]; shade_cell(c1, fill)
    p1 = c1.paragraphs[0]; add_run(p1, desc, size=10, color=MID)
    p1.paragraph_format.space_before = Pt(5); p1.paragraph_format.space_after = Pt(5)
doc.add_paragraph()
body("The platform is currently live and publicly accessible. STPI representatives are invited to explore the full functionality at gcc-playbook.pithonix.ai prior to the first formal meeting. A dedicated demo session, including a live blueprint generation and GOT Advisor demonstration, can be arranged at STPI's convenience.")
divider()

# ── S04: WHY STPI ─────────────────────────────────────────
sec_hdr('04', 'Section Four', 'Why STPI: The Strategic Fit')
body("We did not approach STPI because it is the largest body or the most visible. We approached STPI because, when you look at what GCC investors actually need, and what STPI is actually built to do, the fit is exact. Not approximate. Exact.")
bullet("STPI's statutory mandate covers the registration, support, and monitoring of software technology parks and IT/ITES export units. GCCs are the highest-value, highest-growth category within this mandate. By extending STPI's digital presence to encompass pre-registration GCC discovery, the platform closes the full investor journey, from initial curiosity through to STPI scheme registration, within a single institutional ecosystem.", bold_pre="STPI is the natural institutional home for GCC intelligence.")
bullet("STPI operates 64 centers across India, spanning every major metro and an extensive network of Tier 2 cities. The GCC Playbook's city intelligence database covers 25 cities, encompassing all cities where STPI centers are present. This alignment means that every city recommendation generated by the platform maps directly to a STPI center capable of providing on-the-ground institutional support.", bold_pre="STPI's geographic footprint mirrors our data coverage exactly.")
bullet("Foreign companies visiting stpi.in are not casual browsers. They are organizations that have already identified India as a potential GCC destination and are seeking formal institutional guidance. This is precisely the audience for whom the GCC Playbook delivers maximum value, and precisely the moment at which embedding a simulator produces the highest conversion to formal inquiry.", bold_pre="STPI captures investors at peak intent.")
bullet("In the international investor community, a government institutional endorsement from a MeitY body carries weight that no private-sector marketing campaign can replicate. STPI's co-branding on the GCC Playbook platform would immediately distinguish it from the ecosystem of commercial GCC advisory tools and position it as India's authoritative, government-backed GCC guidance platform.", bold_pre="Co-branding with STPI delivers instant government credibility.")
bullet("Invest India, Make in India, DPIIT, and state investment promotion boards have not yet deployed a personalized, AI-driven investment advisory tool. This partnership positions STPI, under MeitY's leadership, as the first government body in India to offer foreign investors a real-time, AI-generated investment blueprint. This is a significant strategic differentiation.", bold_pre="STPI becomes India's first AI-powered investment promotion authority.")
callout("This is the opportunity to be where it gets started. To move from registration authority to investment generation engine, without changing a single thing about how STPI operates today.",
    bold_pre="STPI has always been where India's technology story gets registered.", gold=True)
divider()
doc.add_page_break()

# ── S05: PARTNERSHIP OPTIONS ──────────────────────────────
sec_hdr('05', 'Section Five', 'Partnership Options: Three Models')
body("Pithonix AI proposes three distinct partnership structures, each calibrated to different levels of STPI institutional commitment, procurement timeline, and strategic ambition. The models are not mutually exclusive. They represent a logical progression, with Model A as the immediate first step.")

for title, fill, is_rec, items in [
    ('MODEL A  |  EMBED  (Recommended)', 'E8F0FA', True, [
        'STPI embeds GCC Playbook as a widget or iframe on stpi.in/gcc or gcc.stpi.in subdomain',
        'Zero STPI engineering effort. Pithonix maintains platform fully',
        'STPI + Pithonix co-branding on platform header',
        'Leads flow to both STPI and Pithonix in real time',
        '10 to 15% revenue share to STPI on GCC engagements originating from STPI traffic',
        'Live within 2 weeks of Letter of Intent. No procurement cycle required',
    ]),
    ('MODEL B  |  WHITE LABEL', 'F5F8FD', False, [
        'STPI deploys a fully STPI-branded version on a custom domain with STPI visual theme',
        'STPI owns the user-facing product entirely',
        "Pithonix provides AI engine and data layer under a technology license agreement",
        'Annual technology license fee, to be agreed based on usage tier',
        'Requires STPI IT procurement process (4 to 8 weeks)',
    ]),
    ('MODEL C  |  CO-DEVELOPMENT', 'F5F8FD', False, [
        "Joint platform: STPI + Pithonix build an expanded version incorporating STPI's proprietary data",
        'STPI contributes registered unit database, park vacancy data, and scheme disbursement records',
        'Pithonix contributes full AI stack, GOT reasoning engine, and market benchmark data',
        'Joint IP ownership with agreed commercialization terms',
        'Funded through MeitY Digital India / STPI Innovation Fund. 12 to 18 month timeline',
    ]),
]:
    p = doc.add_paragraph(); shade_para(p, fill)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, size=12, color=GOLD if is_rec else NAVY, font='Arial')
    for item in items:
        pi = doc.add_paragraph(); shade_para(pi, fill)
        pi.paragraph_format.space_before = Pt(1); pi.paragraph_format.space_after = Pt(1)
        pi.paragraph_format.left_indent = Cm(0.6)
        add_run(pi, '\u2022  ' + item, size=10, color=MID)

doc.add_paragraph()
subhead('Comparative Model Summary')
make_table(
    ['Dimension', 'Model A: Embed', 'Model B: White Label', 'Model C: Co-Dev'],
    [
        ['Time to Live',            '2 weeks',           '4 to 8 weeks',       '12 to 18 months'],
        ['STPI Engineering Effort', 'Zero',              'Minimal',            'Significant'],
        ['STPI Cost',               'Zero',              'Annual license fee',  'Co-investment required'],
        ['Brand Ownership',         'Co-branded',        'STPI brand primary', 'Joint IP'],
        ['Lead Data Access',        'Shared (real-time)','STPI-owned',         'Joint ownership'],
        ['Procurement Required',    'No',                'Yes (standard IT)',  'Yes (formal tender)'],
        ['Revenue to STPI',         '10 to 15% royalty', 'Platform ownership', 'Joint commercialization'],
        ['Pithonix Recommendation', '>>Start Here',      'Upgrade Path',       'Strategic Phase 2'],
    ],
    widths=[1.9, 1.5, 1.6, 1.5]
)
doc.add_paragraph()
callout("It requires no procurement cycle, no STPI engineering resource, no IT infrastructure change, and can be deployed live within two weeks of a signed Letter of Intent. Once 90-day traction data is available, the institutional case for upgrading to Model B or initiating a co-development conversation becomes straightforward to make internally.",
    bold_pre="Our firm recommendation is to begin with Model A.")
divider()

# ── S06: TECHNICAL BRIEF ──────────────────────────────────
sec_hdr('06', 'Section Six', 'Technical Brief: Integration and Data Security')
body("Technical integration under Model A is deliberately designed to be the simplest possible engineering operation. The following details address the integration effort, data security posture, and hosting flexibility in full.")
bullet("Model A integration requires STPI's web team to insert a single iframe embed tag or JavaScript snippet into the stpi.in CMS, an operation functionally identical to embedding a YouTube video or a Google Map. No API keys, no server-side changes, no database modifications, no security review cycle. Estimated effort: under 30 minutes.", bold_pre="Integration effort for STPI: one line of code.")
bullet("The GCC Simulator is architecturally stateless. User inputs are sent to the Gemini API via a secure server-side proxy and the AI response is returned directly to the user's browser. No user data is written to any Pithonix database. There are no user accounts, no cookies beyond session state, and no personally identifiable information collected unless the user voluntarily submits an inquiry form.", bold_pre="Stateless architecture: no user data stored.")
bullet("The Gemini API key is stored and accessed exclusively on the server-side proxy (a Vercel serverless function). It is never transmitted to or accessible from the user's browser at any point. This architecture meets standard government API security requirements and is consistent with MeitY's cloud security guidelines.", bold_pre="API security: server-side key management.")
bullet("If STPI requires data residency within India or mandates hosting on government-approved infrastructure, the platform can be fully mirrored on NIC Cloud or MeghRaj within four weeks. Pithonix will bear the cost of migration under Model A.", bold_pre="Government cloud hosting available within 4 weeks.")
bullet("The current Vercel deployment provides 99.99% uptime with global CDN delivery. Average page load time is under 1.2 seconds on a standard broadband connection and under 4 seconds on a 3G mobile connection.", bold_pre="Uptime and performance.")
bullet("The platform is WCAG 2.1 AA compliant, mobile-responsive across all major device categories, and tested across Chrome, Firefox, Edge, and Safari.", bold_pre="Accessibility and compliance.")
callout("Pithonix AI is prepared to provide STPI's IT team with a complete technical architecture document, security review package, and third-party penetration test report (if required) as part of the due diligence process prior to Model A deployment. We expect no barriers to clearance.")
divider()
doc.add_page_break()

# ── S07: CO-BRANDING ──────────────────────────────────────
sec_hdr('07', 'Section Seven', 'Co-Branding Proposal')
body("The co-branding design philosophy for this partnership is one of institutional elevation: STPI's visual identity leads, with Pithonix AI positioned as the technology provider in the background. This ensures that foreign investors experience STPI as the authoritative source and Pithonix as the enabling technology.")
subhead('Platform Header: Co-Branding Mockup')
t = doc.add_table(rows=2, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
shade_cell(t.rows[0].cells[0], '003F87'); shade_cell(t.rows[0].cells[1], '003F87')
p0 = t.rows[0].cells[0].paragraphs[0]
add_run(p0, 'STPI: Software Technology Parks of India', bold=True, size=11, color=WHITE, font='Arial')
p0.paragraph_format.space_before = Pt(8); p0.paragraph_format.space_after = Pt(8)
p1 = t.rows[0].cells[1].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p1, 'GCC Playbook  |  Powered by Pithonix AI', size=10, color=GOLD, font='Arial')
p1.paragraph_format.space_before = Pt(8); p1.paragraph_format.space_after = Pt(8)
shade_cell(t.rows[1].cells[0], 'E8F0FA')
t.rows[1].cells[0].merge(t.rows[1].cells[1])
pt = t.rows[1].cells[0].paragraphs[0]
add_run(pt, 'An initiative supported by STPI, Ministry of Electronics & Information Technology (MeitY)', italic=True, size=10, color=NAVY)
pt.paragraph_format.space_before = Pt(6); pt.paragraph_format.space_after = Pt(6)
doc.add_paragraph()
bullet("STPI logo and wordmark on the left (primary, larger). 'Powered by Pithonix AI' on the right (secondary, smaller). This follows standard government-endorsed technology platform conventions, as seen on Aarogya Setu, UMANG, and other MeitY-supported digital initiatives.", bold_pre="Dual-logo header architecture.")
bullet("STPI's brand blue (#003F87) will be the primary UI color for all navigation, section headings, and button elements. STPI's green (#00A651) will be applied to positive indicators and success states. Pithonix's gold accent (#C8963E) is retained only for the 'Powered by' attribution elements.", bold_pre="STPI color palette applied platform-wide.")
bullet("Every page of the platform, when accessed via STPI's domain, will display the tagline: 'An initiative supported by STPI, Ministry of Electronics & Information Technology, Government of India.'", bold_pre="Government endorsement tagline on all pages.")
bullet("A future enhancement, planned for the 90-day post-launch development cycle, will add the STPI official seal to the downloadable GCC Blueprint PDF, alongside the STPI center contact details most relevant to the investor's recommended city.", bold_pre="STPI seal on blueprint output (future feature).")
bullet("The Government Incentive Cards module will be re-ordered when accessed via STPI's domain to feature the STPI STP Scheme and STPI Software Technology Park benefits as the first and most prominent recommendations for eligible entities.", bold_pre="STPI scheme as primary incentive recommendation.")
bullet("The 25 cities currently in the intelligence database will be augmented, in coordination with STPI's program team, to reflect STPI's full 64-center geographic footprint, including STPI center contact details and STPI park-specific incentive information.", bold_pre="All 64 STPI center cities reflected in city intelligence.")
body("All co-branding specifications are subject to final review and approval by STPI's communications and branding team. Pithonix will provide full design mockups in Figma format for STPI's review within 5 business days of an MOU being signed.")
divider()

# ── S08: IMPACT ───────────────────────────────────────────
sec_hdr('08', 'Section Eight', 'Impact Projections')
body("The following impact projections are derived from three input assumptions: (1) STPI portal traffic estimated at 500,000+ visits per month based on Similarweb data for stpi.in; (2) GCC Playbook engagement rates benchmarked against comparable B2B AI tool conversion data; and (3) GCC establishment conversion rates based on NASSCOM's 2024 data on qualified investor-to-GCC conversion. These are conservative estimates, presented for planning purposes, not as guaranteed outcomes.")
make_table(
    ['Impact Metric', 'Year 1', 'Year 2', 'Basis'],
    [
        ['GCC Simulator Runs',              '5,000+',         '20,000+',          '1% of STPI monthly traffic x 12 months, growing with awareness'],
        ['Qualified Investor Inquiries',    '300+',           '1,200+',           '6% of simulator runs converting to formal inquiry submission'],
        ['GCC Engagements Influenced',      '20 to 30',       '80 to 120',        '7 to 10% of qualified inquiries advancing to GCC engagement'],
        ['GCC FTE Created (Influenced)',    '2,000 to 5,000', '10,000 to 20,000', 'Average GCC size of 100 to 200 FTE at 3-year maturity'],
        ['STPI Registrations Influenced',   '15 to 25',       '60 to 100',        '70 to 80% of GCC engagements selecting STPI STP scheme'],
        ['Incremental STPI Scheme Revenue', 'Indicative',     'Significant',      'Based on standard STPI registration and service fee structure'],
    ],
    widths=[2.0, 1.0, 1.2, 3.3]
)
doc.add_paragraph()
body("To contextualise these projections: NASSCOM's 2024 data indicates approximately 100 to 130 new GCCs are established in India annually. If this platform, embedded on STPI's portal, contributes to even 15 to 25 new GCC formations in Year 1, it would represent a material 12 to 20% uplift on the national annual GCC addition rate attributable to a single digital channel.")
callout("Each GCC that traces its India journey to the STPI-Pithonix platform becomes a proof point for STPI's role as an active investment generation authority, not merely a compliance registration body. The long-term institutional value of this repositioning is difficult to quantify but is, in our view, the most important outcome of this partnership.",
    bold_pre="The strategic value extends beyond the numbers.", gold=True)
divider()
doc.add_page_break()

# ── S09: NEXT STEPS ───────────────────────────────────────
sec_hdr('09', 'Section Nine', 'Next Steps and Implementation Timeline')
body("The following implementation timeline applies to Model A (Embed), the recommended starting point. The entire sequence from signed Letter of Intent to public beta launch on stpi.in can be completed within six weeks. The timeline is intentionally aggressive: we have designed Model A to require minimal institutional process on STPI's side, so that momentum is not lost in procurement cycles.")
make_table(
    ['Action / Milestone', 'Responsible Party', 'Timeline', 'Deliverable'],
    [
        ['Proposal review and internal discussion',         'STPI Leadership',      'Week 1 to 2', 'Go / No-Go decision with internal stakeholders'],
        ['Initial meeting and live platform demonstration', 'STPI + Pithonix',      'Week 1',      'Live demo of GCC Simulator, GOT Advisor, and City Intelligence'],
        ['Partnership model selection (A / B / C)',         'STPI + Pithonix',      'Week 2',      'Confirmed model with any customization requirements documented'],
        ['MOU / Letter of Intent drafting and execution',   'Both parties (legal)', 'Week 3 to 4', 'Signed MOU covering co-branding rights, lead data sharing, and revenue share terms'],
        ['Technical integration: Model A embed',            'Pithonix (primary)',   'Week 3 to 4', 'Embed-ready code package delivered to STPI web team for review'],
        ['STPI brand theme applied to platform',            'Pithonix',             'Week 4',      'Platform updated with STPI color palette, co-branding header, and endorsement tagline'],
        ['STPI IT and security review',                     'STPI IT team',         'Week 4 to 5', 'Security clearance for embed code deployment on stpi.in'],
        ['Beta launch on stpi.in',                          'Both parties',         'Week 5 to 6', 'Platform live on stpi.in/gcc or gcc.stpi.in, internal team testing phase'],
        ['Press release and public announcement',           'Both parties (comms)', 'Week 6',      'Joint press release to business media, investor communities, and NASSCOM channels'],
    ],
    widths=[2.2, 1.3, 0.9, 3.1]
)
doc.add_paragraph()
subhead('Immediate Next Step: Requested from STPI')
body("Pithonix AI requests the following as the immediate next action from STPI upon review of this proposal:")
for num, title, detail in [
    ('1', 'Schedule a Live Platform Demonstration',
     'A 45-minute session with the relevant STPI team (DG office, IT cell, or GCC programs team) to walk through the live platform, generate a sample blueprint, and answer technical and commercial questions. Can be in-person (Bengaluru or New Delhi STPI offices) or via video conference.'),
    ('2', 'Designate a Nodal Officer for this Engagement',
     "A named STPI point of contact at Director or Joint Director level to own the partnership evaluation process and coordinate with Pithonix AI's engagement lead."),
    ('3', 'Confirm Interest in Proceeding (Letter of Intent)',
     "Even a one-page non-binding expression of interest from STPI, confirming willingness to explore Model A, enables Pithonix to begin technical customization immediately, compressing the timeline to launch."),
]:
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c0 = t.rows[0].cells[0]; c0.width = Inches(0.45); shade_cell(c0, 'C8963E')
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, num, bold=True, size=14, color=WHITE, font='Arial')
    p0.paragraph_format.space_before = Pt(8); p0.paragraph_format.space_after = Pt(6)
    c1 = t.rows[0].cells[1]
    p1 = c1.paragraphs[0]; add_run(p1, title, bold=True, size=11, color=NAVY, font='Arial')
    p1.paragraph_format.space_before = Pt(4); p1.paragraph_format.space_after = Pt(3)
    p2 = c1.add_paragraph(); add_run(p2, detail, size=10, color=MID)
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(8)
callout("No long procurement cycles. No resource commitments. No risk. Just a live platform, a shared ambition, and two weeks to launch. We are ready when STPI is.",
    bold_pre="We have kept Model A deliberately simple, because we know how government institutions work, and we respect that.")
divider()

# ── S10: ABOUT PITHONIX ───────────────────────────────────
sec_hdr('10', 'Section Ten', 'About PITHONIX AI INDIA PRIVATE LIMITED')
body("We started with a simple belief: that AI should not just assist people at work, it should think alongside them. Not a chatbot. Not a dashboard. A reasoning partner that understands context, connects the dots across domains, and surfaces what nobody thought to ask. That belief became PITHONIX AI INDIA PRIVATE LIMITED, incorporated in Hyderabad on 18th March 2026, DPIIT recognised, and already live with products that prove the concept.")
for icon_lbl, title, text in [
    ('Flagship Product', 'JEET ERP',
     "India's first AI-first enterprise platform, a fully integrated ERP system with 25 specialized AI agents spanning HR, finance, procurement, operations, and customer intelligence. Designed for GCC-stage organizations that require enterprise-grade intelligence without legacy system complexity."),
    ('Core Technology', 'GOT: Graph of Thought',
     "Pithonix's proprietary reasoning architecture. Unlike standard LLM chain-of-thought approaches, GOT constructs a multi-dimensional reasoning graph across business domains, enabling AI agents to navigate complex, multi-variable enterprise decisions with structured, auditable reasoning paths."),
    ('Key System', 'HARI: Human Augmented Realistic Intelligence',
     "Pithonix's structural framework defining how AI agents and humans augment each other. HARI ensures every AI-driven decision, across talent, compliance, finance, and operations, is paired with human context and oversight. No black-box automation: every agent output is human-reviewable, explainable, and auditable."),
    ('Operational System', 'BOT: Business Operations Terminal',
     "A unified AI command layer for business operations, integrating procurement, vendor management, compliance tracking, and reporting into a single conversational interface. Designed to replace the fragmented tool sprawl that characterises GCC operations in their early phase."),
]:
    p_i = doc.add_paragraph(); shade_para(p_i, 'FDF6E8')
    p_i.paragraph_format.space_before = Pt(10); p_i.paragraph_format.space_after = Pt(1)
    add_run(p_i, icon_lbl.upper(), size=8, color=GOLD, font='Arial')
    p_t = doc.add_paragraph(); shade_para(p_t, 'FDF6E8')
    p_t.paragraph_format.space_before = Pt(1); p_t.paragraph_format.space_after = Pt(3)
    add_run(p_t, title, bold=True, size=14, color=NAVY, font='Georgia')
    p_tx = doc.add_paragraph(); shade_para(p_tx, 'FDF6E8')
    p_tx.paragraph_format.space_before = Pt(0); p_tx.paragraph_format.space_after = Pt(10)
    p_tx.paragraph_format.left_indent = Cm(0.3)
    add_run(p_tx, text, size=10, color=MID)

subhead('GCC Practice: End-to-End Capability')
body("Pithonix AI's GCC practice specialises in supporting foreign companies through the complete GCC establishment journey in India, from initial feasibility evaluation through to operational steady-state.")
bullet("Private Limited company incorporation, branch office setup, liaison office registration, STPI scheme registration, DPIIT startup recognition, and all statutory compliance setup under Companies Act 2013.", bold_pre="Entity Formation and Legal Structure.")
bullet("Leveraging the same 25-city intelligence database that powers the GCC Playbook, Pithonix provides clients with personalized location recommendations, SEZ / IT park shortlisting, office space negotiations, and fit-out project management.", bold_pre="Location Selection and Real Estate Advisory.")
bullet("End-to-end talent acquisition for GCC founding team hires, from CTO/CHRO-level leadership to technology individual contributors. HARI AI provides ongoing workforce analytics once the GCC is operational.", bold_pre="Talent Acquisition and HR Systems.")
bullet("Deployment of JEET ERP and integrated AI agent stack tailored to the GCC's specific function mix, ensuring technology infrastructure is in place from Day 1, not retrofitted 18 months into operations.", bold_pre="Technology Deployment.")
bullet("Post-establishment managed services covering payroll compliance, statutory filings, vendor management, and performance reporting, enabling the GCC's India leadership team to focus on capability delivery rather than administrative overhead.", bold_pre="Ongoing Operations Support.")

doc.add_paragraph()
t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
c0 = t.rows[0].cells[0]; c1 = t.rows[0].cells[1]
shade_cell(c0, 'E8F0FA'); shade_cell(c1, 'F5F8FD')
cell_border(c0, sz='8'); cell_border(c1, sz='8')

def contact_head(cell, text):
    p = cell.paragraphs[0]; add_run(p, text.upper(), size=9, color=GOLD, font='Arial')
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)

def contact_line(cell, label, value):
    p = cell.add_paragraph()
    add_run(p, label + ': ', bold=True, size=10, color=NAVY, font='Arial')
    add_run(p, value, size=10, color=DARK)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

contact_head(c0, 'Primary Contact: Partnership Enquiry')
contact_line(c0, 'Organization', 'PITHONIX AI INDIA PRIVATE LIMITED')
contact_line(c0, 'Email', 'info@pithonix.ai')
contact_line(c0, 'Platform', 'gcc-playbook.pithonix.ai')
contact_line(c0, 'Website', 'pithonix.ai')
contact_line(c0, 'Registered', 'Hyderabad, Telangana, India')
contact_line(c0, 'Incorporated', '18th March 2026 | Companies Act, 2013')
contact_line(c0, 'CIN', 'U62090TS2026PTC213220')
contact_line(c0, 'PAN', 'AAQCP8532M  |  TAN: HYDP24088B')

contact_head(c1, 'Platform Access and Demonstration')
contact_line(c1, 'Live Platform', 'gcc-playbook.pithonix.ai')
contact_line(c1, 'Demo Request', "Available in person (Bengaluru / New Delhi / Hyderabad) or via video conference at STPI's convenience.")
contact_line(c1, 'For STPI', 'Please direct all communication to info@pithonix.ai with subject line "STPI GCC Playbook Partnership Enquiry"')

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Confidential  |  Prepared by PITHONIX AI INDIA PRIVATE LIMITED for STPI  |  April 2026  |  gcc-playbook.pithonix.ai', size=8, color=LIGHT, font='Arial')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
add_run(p, 'PITHONIX AI  x  STPI', bold=True, size=11, color=NAVY, font='Arial')
p = doc.add_paragraph(); shade_para(p, 'C8963E')
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(0); p.add_run(' ')

out = 'C:/Users/Administrator/Downloads/gcc-playbook/Pithonix-STPI-Proposal-2026.docx'
doc.save(out)
print('Saved:', out)
'''

with open('generate_proposal_docx.py', 'w', encoding='utf-8') as f:
    f.write(script.lstrip('\n'))

print("Written successfully.")
